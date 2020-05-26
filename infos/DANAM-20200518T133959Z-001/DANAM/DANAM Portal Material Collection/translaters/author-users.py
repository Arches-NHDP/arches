import docx
import xlrd

doc = docx.Document()
file = "NHDP actors und agents part 1.xlsx"

wb = xlrd.open_workbook(file)
g_sheet = wb.sheet_by_name("Sheet1")
for i in range(g_sheet.nrows-1):
    heading = doc.add_paragraph('Monument Name: ' + g_sheet.cell_value(i+1, 1) + '\n')
    #heading.bold = True
    id = heading.add_run('Identifier: ' + g_sheet.cell_value(i+1, 0) + '\n')
    #id.bold = True
    detail = doc.add_paragraph('Editor(s): ' + g_sheet.cell_value(i+1, 2) + '\n')
    detail.add_run('Field work, collection of (art) historical, religious and anthropological data: '
                   + g_sheet.cell_value(i+1, 3) + '\n')
    detail.add_run('Drawings and architectural data: ' + g_sheet.cell_value(i + 1, 6) + '\n')
    detail.add_run('Photography after 2015: ' + g_sheet.cell_value(i + 1, 5) + '\n')
    detail.add_run('History and revision: ' + g_sheet.cell_value(i + 1, 4) + '\n')
    detail.add_run('Visual data management: Matthias Arnold, Lizeth Ortiz \n')
    detail.add_run('Copy editing: ' + g_sheet.cell_value(i + 1, 7) + '\n')
    doc.add_page_break()
doc.save("NHDP authors and agents part 1.docx")